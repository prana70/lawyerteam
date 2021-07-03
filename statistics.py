#coding:utf-8

import pandas as pd
import datetime
import time
import os
from sqlalchemy import create_engine

from os import startfile
from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches,Pt,Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_UNDERLINE,WD_TAB_ALIGNMENT,WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

import unprojected as up
import re

import warnings

warnings.filterwarnings('ignore')


def MakeFullWorklog(start_date,end_date): #生成指定起止期间的归属项目库的工作日志DataFrame
    start_time=time.clock()
    engine=create_engine('sqlite:///data.db')
    worklog_df=pd.read_sql('worklog',engine)
    worklog_df=worklog_df[(worklog_df.办理日期>=start_date)&(worklog_df.办理日期<=end_date)]
    print('初始记录数：',len(worklog_df))

    project_df=pd.read_excel(os.getcwd()+'\\config\\项目库.xlsx')

    #print(worklog_df.dtypes)
    #print(project_df.dtypes)

    #从律师角度统计
    #print(df.groupby(['填报人',df.办理日期.dt.to_period('D')])['耗费时间'].sum().loc['曾婕'])

    #从客户角度统计
    w_df=pd.DataFrame()
    for index,row in project_df.iterrows():
        customer_expression=up.MakeRegularExpression(row.客户名称关键字)
        project_expression=up.MakeRegularExpression(row.项目名称关键字)
        pw_df=worklog_df[(worklog_df.客户名称.str.match(customer_expression)) &
                         (worklog_df.项目名称.str.match(project_expression)) &
                         (worklog_df.办理日期.dt.date>=row.起始日期.date()) &
                         (worklog_df.办理日期.dt.date<=row.终止日期.date())]
        worklog_df=worklog_df.drop(pw_df.index)
        #print(len(worklog_df))
        pw_df['项目库客户名称']=row.客户名称
        pw_df['项目库项目名称']=row.项目名称
        pw_df['项目库律师费总额']=row.律师费总额
        pw_df['项目库办案律师费']=row.办案律师费
        w_df=w_df.append(pw_df)

    worklog_df['项目库客户名称']='未见'
    worklog_df['项目库项目名称']='未见'
    worklog_df['项目库律师费总额']=0
    worklog_df['项目库办案律师费']=0
    print('未入库记录数：',len(worklog_df))
    w_df=w_df.append(worklog_df)
    print('最终记录数：',len(w_df))
    #w_df.to_csv('日志归库.csv')
    end_time=time.clock()
    print('耗时：',end_time-start_time)
    return w_df

if __name__=='__main__':
    
    customer_name='每经传媒' #input('请输入客户名称：')

    report_method=input('请输入出报告的方式：\n  0、按周出报告\n  1、按自定义期间出报告\n  ')
    report_unit={'0':'周','1':'期'}
    report_subtitle={'0':'一周法律事务','1':'期间法律事务'}



    if report_method=='0':
        current_week_no=time.strftime('%W')
        current_year=time.strftime('%Y')
        year_=input('现在是'+current_year+'年的第'+current_week_no+'周，您打算出周报的年份：')
        week_no=input('第几周的周报：')
        start_date=datetime.datetime.strptime(year_+'-W'+week_no+'-1','%Y-W%W-%w').date()
        end_date=datetime.datetime.strptime(year_+'-W'+week_no+'-0','%Y-W%W-%w').date()
        report_no=week_no
    else:
        yy,mm,dd=map(int,re.split(r'[\s\,\.\\:;，。；；]',input('请输入报告起始日期，格式如 2020,11,20：')))
        start_date=datetime.date(yy,mm,dd)
        yy,mm,dd=map(int,re.split(r'[\s\,\.\\:;，。；；]',input('请输入报告截止日期，格式如 2020,11,20：')))
        end_date=datetime.date(yy,mm,dd)
        year_=input('请输入报告的年份：')
        report_no=input('请输入报告的期数：')


    
    df=MakeFullWorklog(start_date,end_date)
    df=df[(df.项目库客户名称==customer_name)]
    df=df[['项目库项目名称','办理日期','事务类型','耗费时间','工作内容']]
    
    doc=Document(os.getcwd()+'\\report\\每经一周法律事务报告空白模版.docx')

    #添加标题行
    doc.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph0_run=doc.paragraphs[0].add_run('法律工作报告')
    paragraph0_run.font.color.rgb=RGBColor(255,0,0)
    paragraph0_run.font.size=Pt(26)
    paragraph0_run.font.name='方正小标宋_GBK'
    paragraph0_run._element.rPr.rFonts.set(qn('w:eastAsia'),'方正小标宋_GBK')
    doc.paragraphs[0].style.paragraph_format.line_spacing=1
    doc.paragraphs[0].style.paragraph_format.space_after=Cm(0)

    #添加第二行（空行）
    paragraph1=doc.add_paragraph()

    #添加第三行
    paragraph2=doc.add_paragraph()
    paragraph2.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph2_run=paragraph2.add_run('〔'+year_+'〕'+report_no+report_unit[report_method]+'        '+str(datetime.datetime.now().strftime('%Y年%m月%d日')))
    paragraph2_run.font.color.rgb=RGBColor(255,0,0)
    paragraph2_run.font.size=Pt(16)
    paragraph2_run.font.name='方正小标宋_GBK'
    paragraph2_run._element.rPr.rFonts.set(qn('w:easeAsia'),'方正小标宋_GBK')

    #添加第四行
    paragraph3=doc.add_paragraph()
    paragraph3.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph3_run0=paragraph3.add_run(report_subtitle[report_method])
    paragraph3_run0.font.size=Pt(20)
    paragraph3_run0.font.name='方正小标宋_GBK'
    paragraph3_run0._element.rPr.rFonts.set(qn('w:eastAsia'),'方正小标宋_GBK')

    paragraph3_run1=paragraph3.add_run('（'+str(start_date.strftime('%Y.%m.%d'))+'-'+str(end_date.strftime('%Y.%m.%d'))+'）')
    paragraph3_run1.font.size=Pt(16)
    paragraph3_run1.font.name='方正小标宋_GBK'
    paragraph3_run1._element.rPr.rFonts.set(qn('w:eastAsia'),'方正小标宋_GBK')

    #添加第五行（空行）
    paragraph4=doc.add_paragraph()
    paragraph4.style=doc.styles['Normal']

    #设置style-H1、N1
    doc.styles.add_style('H1',WD_STYLE_TYPE.PARAGRAPH)
    doc.styles['H1'].font.name='方正黑体_GBK'
    doc.styles['H1']._element.rPr.rFonts.set(qn('w:eastAsia'),'方正黑体_GBK')
    doc.styles['H1'].font.size=Pt(16)
    doc.styles['H1'].font.bold=True
    doc.styles['H1'].paragraph_format.first_line_indent=Cm(1.1)
    doc.styles['H1'].paragraph_format.line_spacing=Cm(1)

    doc.styles.add_style('N1',WD_STYLE_TYPE.PARAGRAPH)
    doc.styles['N1'].font.name='方正仿宋_GBK'
    doc.styles['N1']._element.rPr.rFonts.set(qn('w:eastAsia'),'方正仿宋_GBK')
    doc.styles['N1'].font.size=Pt(16)
    doc.styles['N1'].paragraph_format.first_line_indent=Cm(1.2)
    doc.styles['N1'].paragraph_format.line_spacing=1

    #添加（一、诉讼事务）
    paragraph_lawsuit=doc.add_paragraph('一、诉讼事务',style='H1')
    df_lawsuit=df[df.项目库项目名称.str.contains('VS|Vs|vS|vs')] #从工作日志中筛选出诉讼类工作日志
    if len(df_lawsuit)>0:
        df=df.drop(df_lawsuit.index) #从工作日志中删除已筛选出的诉讼类工作日志
        group_lawsuit=df_lawsuit.groupby(by=['项目库项目名称']) #按诉讼项目分组
        #每一分组填充一段内容
        for key,df_inter in group_lawsuit:
            text=key+'：'
            i=len(list(df_inter.iterrows()))
            for row in df_inter.iterrows():
                text+=row[1]['办理日期'].date().strftime('%Y年%m月%d日')+'，'+row[1]['工作内容']
                i-=1
                if i>0:
                    text+='；'
                else:
                    text+='。'
            doc.add_paragraph(text,style='N1')
    else:
        doc.add_paragraph('无。',style='N1')

    #添加（二、经营事务）
    paragraph_operate=doc.add_paragraph('二、经营事务',style='H1')
    df_consult=df[df.项目库项目名称.str.contains('法律顾问')] #从工作日志中筛选出法律顾问类工作日志，用于采编事务和综合事务的内容
    df_nolawsuit=df.drop(df_consult.index) #从工作日志中删除法律顾问类工作日志，剩下为非诉专项工作日志，用于经营事务的内容。（诉讼类在之前已删除）
    
    if len(df_nolawsuit)>0:
        group_nolawsuit=df_nolawsuit.groupby(by=['项目库项目名称'])
        for key,df_inter in group_nolawsuit:
            text=key+'：'
            i=len(list(df_inter.iterrows()))
            for row in df_inter.iterrows():
                text+=row[1]['办理日期'].date().strftime('%Y年%m月%d日')+'，'+row[1]['工作内容']
                i-=1
                if i>0:
                    text+='；'
                else:
                    text+='。'
            doc.add_paragraph(text,style='N1')
    else:
        doc.add_paragraph('无。',style='N1')


    #添加（三、采编事务）
    paragraph_edit=doc.add_paragraph('三、采编事务',style='H1')
    df_edit=df_consult[df_consult.事务类型.str.contains('意见报告')] #从法律顾问类工作日志中筛选出报告意见类工作日志，用于采编事务的内容
    df_consult=df_consult.drop(df_edit.index) #从法律顾问类工作日志中删除报告意见类工作日志，剩下工作日志用于综合事务的内容
    
    if len(df_edit)>0:
        for row in df_edit.iterrows():
            text=row[1]['办理日期'].date().strftime('%Y年%m月%d日')+'，'+row[1]['工作内容']+'。'
            doc.add_paragraph(text,style='N1')
    else:
        doc.add_paragraph('无。',style='N1')

    #添加（四、综合事务）
    paragraph_consult=doc.add_paragraph('四、综合事务',style='H1')
    event_unit={'合同审撰':'份','文件审撰':'份','咨询沟通':'次','发律师函':'次','事实分析':'次','其他':'件','文件传递':'次','文件整理':'次','法律培训':'次'}

    if len(df_consult)>0:
        group_consult=df_consult.groupby(by=['事务类型'])['事务类型'].count()
        text='受相关部门委托，完成成都公司、北京公司、上海分公司、上海经闻公司'
        i=len(group_consult.index)
        for ele in group_consult.index:
            print(ele,group_consult[ele])
            text+=ele+str(group_consult[ele])+event_unit[ele]
            i-=1
            if i>0:
                text+='，'
            else:
                text+='。'
        doc.add_paragraph(text,style='N1')
    else:
        doc.add_paragraph('无。',style='N1')


    #添加（完）
    paragraph_end=doc.add_paragraph('（完）',style='N1')
    paragraph_end.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.save(os.getcwd()+'\\report\\'+customer_name+year_+'年第'+report_no+report_unit[report_method]+'法律事务报告.docx')
    startfile((os.getcwd()+'\\report\\'+customer_name+year_+'年第'+report_no+report_unit[report_method]+'法律事务报告.docx'))
    
